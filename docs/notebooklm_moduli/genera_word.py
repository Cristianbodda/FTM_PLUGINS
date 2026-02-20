"""
Genera il documento Word con tabella moduli e passi da seguire
per NotebookLM Audio Overview -> Video Corso
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Stili
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# === COPERTINA ===
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FTM Academy\nVideo Corso Coach/Formatore')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Guida alla Produzione con NotebookLM')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x76, 0x4B, 0xA2)

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Versione 1.0 - Febbraio 2026\nFTM Academy - Gestione Competenze Professionali')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# === INDICE ===
doc.add_heading('Indice', level=1)
doc.add_paragraph('1. Panoramica del Progetto')
doc.add_paragraph('2. Tabella dei 10 Moduli')
doc.add_paragraph('3. Passi per la Produzione')
doc.add_paragraph('4. Prompt Personalizzati per NotebookLM')
doc.add_paragraph('5. Checklist Finale')

doc.add_page_break()

# === SEZIONE 1: PANORAMICA ===
doc.add_heading('1. Panoramica del Progetto', level=1)

doc.add_paragraph(
    'Questo documento descrive il processo per creare un video corso completo '
    'per coach e formatori FTM Academy, utilizzando Google NotebookLM per generare '
    'l\'audio narrato e PowerPoint per le slide visive.'
)

doc.add_heading('Flusso di Produzione', level=2)
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow.add_run(
    'Manuale (3603 righe .md)\n'
    '          \u2193\n'
    '10 Documenti per Modulo\n'
    '          \u2193\n'
    'Carica su NotebookLM (1 notebook per modulo)\n'
    '          \u2193\n'
    'Genera Audio Overview per ogni modulo\n'
    '          \u2193\n'
    'Scarica audio (.wav/.mp3)\n'
    '          \u2193\n'
    'Abbina a slide PowerPoint\n'
    '          \u2193\n'
    'VIDEO CORSO COMPLETO'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# === SEZIONE 2: TABELLA MODULI ===
doc.add_heading('2. Tabella dei 10 Moduli', level=1)

doc.add_paragraph(
    'Ogni modulo corrisponde a un file .md separato, '
    'pronto per essere caricato su NotebookLM.'
)

# Tabella
moduli = [
    ('1', 'Modulo_01_Introduzione_Accesso.md', 'Cap. 1-2', 'Introduzione e Accesso al Sistema',
     "Cos'\u00e8 FTM Academy, i 5 strumenti del coach, le 4 fonti dati, flusso di lavoro, accesso al sistema", '~890'),
    ('2', 'Modulo_02_Dashboard_Viste_Filtri.md', 'Cap. 3-5', 'Coach Dashboard - Panoramica, Viste e Filtri',
     'Struttura dashboard, pulsanti header, 5 statistiche, 6 quick filters, 4 viste, zoom, filtri avanzati', '~1988'),
    ('3', 'Modulo_03_Card_Studente_Azioni.md', 'Cap. 6', 'Card Studente e Azioni Rapide',
     'Card studente, header/badge, 3 barre progresso, badge stato, timeline, atelier, scelte, 7 quick actions', '~1176'),
    ('4', 'Modulo_04_Report_Studente_Tab_Radar.md', 'Cap. 7-8', 'Report Studente - Tab e Radar',
     '7 tab principali, 6 tab FTM, radar SVG/Chart.js, valutazione inline, selettore quiz, stampa personalizzata', '~3678'),
    ('5', 'Modulo_05_Gap_Analysis_Spunti.md', 'Cap. 9', 'Gap Analysis e Spunti Colloquio',
     'Gap Analysis, soglie configurabili, indicatori sopra/sottovalutazione, spunti colloquio, tono formale/colloquiale', '~910'),
    ('6', 'Modulo_06_Valutazione_Scala_Bloom.md', 'Cap. 10 + App.B', 'Valutazione Formatore e Scala Bloom',
     'Pagina valutazione, selezione settore, compilazione aree, scala Bloom 0-6 con esempi per settore, salvataggio stati', '~2403'),
    ('7', 'Modulo_07_Bilancio_Note_Coach.md', 'Cap. 11-12', 'Bilancio Competenze e Note Coach',
     '6 tab bilancio (Panoramica, Radar, Mappa, Confronta, Colloquio, Matching), note coach, visibilit\u00e0', '~2291'),
    ('8', 'Modulo_08_SelfAssessment_Scheduler.md', 'Cap. 13-14', 'Self-Assessment e FTM Scheduler',
     'Dashboard autovalutazioni, azioni studente, Scheduler con calendario, gruppi, aule, atelier, presenze', '~3819'),
    ('9', 'Modulo_09_Casi_Uso_Pratici.md', 'Cap. 15', 'Casi d\'Uso Pratici',
     '16 workflow passo-passo: nuovo studente, colloquio, valutazione, confronto, export, scheduler, filtri', '~1825'),
    ('10', 'Modulo_10_Problemi_Riferimenti.md', 'Cap. 16 + App.A,C,D', 'Risoluzione Problemi e Riferimenti',
     '14 problemi comuni, contatti supporto, glossario, checklist settimanale, mappa navigazione, URL, AJAX', '~2639'),
]

table = doc.add_table(rows=1, cols=6)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
headers = ['#', 'File', 'Capitoli', 'Titolo Modulo', 'Contenuto Principale', 'Parole']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)

# Dati
for m in moduli:
    row = table.add_row()
    for i, val in enumerate(m):
        cell = row.cells[i]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# Larghezze colonne
widths = [Cm(1), Cm(4.5), Cm(2), Cm(3.5), Cm(5.5), Cm(1.5)]
for row in table.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

doc.add_page_break()

# === SEZIONE 3: PASSI PRODUZIONE ===
doc.add_heading('3. Passi per la Produzione', level=1)

steps = [
    ('Passo 1: Preparazione File', [
        'I 10 file .md sono gi\u00e0 pronti nella cartella docs/notebooklm_moduli/',
        'Ogni file ha un header descrittivo che guida NotebookLM sul contenuto',
        'Verifica che tutti i file siano presenti (vedi tabella sopra)',
    ]),
    ('Passo 2: Configurare NotebookLM', [
        'Vai su notebooklm.google.com',
        'Accedi con il tuo account Google',
        'Clicca sull\'icona ingranaggio (Settings) in alto a destra',
        'Vai su "Output Language" e seleziona "Italiano"',
        'Questa impostazione far\u00e0 generare l\'audio in italiano',
    ]),
    ('Passo 3: Creare i Notebook (uno per modulo)', [
        'Clicca "New Notebook" (o "+ Crea")',
        'Rinomina il notebook: es. "FTM Corso - Modulo 1: Introduzione"',
        'Clicca "Add source" > "Upload" e carica il file .md del modulo',
        'Attendi che NotebookLM processi il documento (pochi secondi)',
        'Ripeti per tutti i 10 moduli',
    ]),
    ('Passo 4: Generare l\'Audio Overview', [
        'Nel pannello destro "Studio", trova "Audio Overview"',
        'Clicca su "Customize" (Personalizza) - IMPORTANTE',
        'Inserisci il prompt personalizzato (vedi sezione 4)',
        'Scegli formato: "Single speaker" (lezione) o "Deep dive" (dialogo)',
        'Scegli lunghezza: "Default" (5-10 min) o "Longer" (10-20 min)',
        'Clicca "Generate" e attendi 2-5 minuti',
    ]),
    ('Passo 5: Scaricare gli Audio', [
        'Quando l\'audio \u00e8 pronto, clicca i tre puntini (menu)',
        'Seleziona "Download"',
        'Rinomina il file: es. "Audio_Modulo_01_Introduzione.wav"',
        'Ripeti per tutti i 10 moduli',
    ]),
    ('Passo 6: Preparare i PowerPoint', [
        'I 10 file .pptx sono gi\u00e0 pronti nella cartella docs/notebooklm_moduli/',
        'Ogni presentazione ha slide con placeholder per screenshot',
        'Sostituisci i placeholder con screenshot reali della piattaforma',
        'Aggiungi eventuali slide extra con contenuti aggiuntivi',
    ]),
    ('Passo 7: Assemblare il Video Corso', [
        'Per ogni modulo: apri il PowerPoint corrispondente',
        'Inserisci > Audio > Audio nel PC > seleziona il file audio',
        'Oppure usa un video editor (CapCut, DaVinci, Camtasia)',
        'Sincronizza le slide con l\'audio',
        'Esporta come video (.mp4)',
    ]),
]

for title, items in steps:
    doc.add_heading(title, level=2)
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph('')

doc.add_page_break()

# === SEZIONE 4: PROMPT PERSONALIZZATI ===
doc.add_heading('4. Prompt Personalizzati per NotebookLM', level=1)

doc.add_paragraph(
    'Per ogni modulo, usa il prompt base qui sotto nel campo "Customize" '
    'di NotebookLM. Puoi personalizzarlo aggiungendo la riga specifica per il modulo.'
)

doc.add_heading('Prompt Base (uguale per tutti)', level=2)
prompt_base = doc.add_paragraph()
prompt_base.style = 'No Spacing'
run = prompt_base.add_run(
    'Crea una lezione formativa in italiano per coach/formatori professionali.\n'
    'Il tono deve essere professionale ma accessibile, come un formatore esperto\n'
    'che spiega a un collega come usare il sistema FTM Academy.\n\n'
    'Concentrati su:\n'
    '- Spiegare passo-passo ogni funzionalit\u00e0 descritta\n'
    '- Dare suggerimenti pratici su quando e come usare ogni strumento\n'
    '- Evidenziare gli errori comuni da evitare\n\n'
    'NON usare un formato podcast informale. Questo \u00e8 materiale formativo\n'
    'per un video corso aziendale.'
)
run.font.size = Pt(10)
run.font.name = 'Consolas'

doc.add_heading('Aggiunte Specifiche per Modulo', level=2)

prompt_specifici = [
    ('Modulo 1', 'Dai una panoramica chiara del sistema. Spiega perch\u00e9 esistono 4 fonti dati diverse e come si collegano tra loro. Questo \u00e8 il primo modulo, quindi sii accogliente.'),
    ('Modulo 2', 'Concentrati sulle 4 viste della Dashboard. Spiega quando usare ciascuna vista e perch\u00e9. Descrivi i filtri come strumenti per risparmiare tempo.'),
    ('Modulo 3', 'Descrivi ogni elemento della card studente come se stessi guardando lo schermo insieme al coach. Spiega cosa significano i colori e le icone.'),
    ('Modulo 4', 'Questo \u00e8 il modulo pi\u00f9 denso. Concentrati sui tab pi\u00f9 importanti: Panoramica (radar), Dettagli (tabella) e la valutazione inline. Spiega come leggere un radar.'),
    ('Modulo 5', 'Spiega la Gap Analysis con esempi concreti. Usa numeri: "Se lo studente si auto-valuta 80% ma il quiz dice 40%, il gap \u00e8 +40% sopravvalutazione". Spiega come preparare un colloquio usando questi dati.'),
    ('Modulo 6', 'Questa lezione \u00e8 cruciale. Spiega ogni livello Bloom con gli esempi pratici per meccanica e automobile. Aiuta il coach a distinguere tra "Applicare" e "Analizzare" che \u00e8 l\'errore pi\u00f9 comune.'),
    ('Modulo 7', 'Spiega la differenza tra Report Studente e Bilancio Competenze. Il Bilancio \u00e8 per preparare colloqui, il Report per analizzare dati. Descrivi il workflow del colloquio.'),
    ('Modulo 8', 'Dividi la lezione in due parti: prima Self-Assessment (pi\u00f9 semplice), poi Scheduler (pi\u00f9 complesso). Per lo Scheduler concentrati su: creare gruppo, vista settimanale, registrare presenze.'),
    ('Modulo 9', 'Presenta i casi d\'uso come scenari reali. Usa un tono narrativo: "Immagina che arrivi un nuovo studente luned\u00ec mattina...". Questo modulo deve essere il pi\u00f9 coinvolgente.'),
    ('Modulo 10', 'Per i problemi comuni, sii pratico e rassicurante: "Se ti capita X, niente panico, fai Y". Chiudi con la checklist settimanale come riepilogo dell\'intero corso.'),
]

for title, text in prompt_specifici:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True

doc.add_page_break()

# === SEZIONE 5: CHECKLIST ===
doc.add_heading('5. Checklist Finale', level=1)

checklist = [
    '\u2610 10 file .md creati e verificati',
    '\u2610 NotebookLM configurato in italiano',
    '\u2610 10 notebook creati su NotebookLM',
    '\u2610 10 file caricati (uno per notebook)',
    '\u2610 10 Audio Overview generati con prompt personalizzati',
    '\u2610 10 file audio scaricati e rinominati',
    '\u2610 10 PowerPoint con screenshot reali inseriti',
    '\u2610 10 video assemblati (audio + slide)',
    '\u2610 Video finale esportato in .mp4',
    '\u2610 Test di visione completo effettuato',
]

for item in checklist:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.size = Pt(12)

# Salva
output_path = os.path.join(os.path.dirname(__file__), 'FTM_Video_Corso_Guida_Produzione.docx')
doc.save(output_path)
print(f'Word salvato: {output_path}')
